from __future__ import annotations

import json
import re
import unicodedata
import html as html_lib
from pathlib import Path
from typing import Any, Dict, Iterable, Iterator, List, Optional, Tuple, Union

import docx

try:
    from bs4 import BeautifulSoup
    import mammoth
except Exception:  # pragma: no cover
    BeautifulSoup = None
    mammoth = None


def collapse_ws(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip()


def strip_vietnamese_accents(text: str, keep_dd: bool = False) -> str:
    if keep_dd:
        text = text.replace("đ", "dd").replace("Đ", "dd")
    else:
        text = text.replace("đ", "d").replace("Đ", "D")
    return unicodedata.normalize("NFKD", text or "").encode("ASCII", "ignore").decode("utf-8")


def slugify(text: str) -> str:
    text = strip_vietnamese_accents(text or "", keep_dd=True)
    text = text.replace("/", "_").replace("-", "_")
    text = re.sub(r"\s+", "_", text).lower()
    text = re.sub(r"[^a-z0-9_\.]", "", text)
    text = re.sub(r"_+", "_", text).strip("_")
    return text or "unknown"


def safe_dirname(text: str) -> str:
    text = strip_vietnamese_accents(text or "unknown", keep_dd=False)
    text = re.sub(r"[\\/*?:\"<>|]", "_", text)
    text = re.sub(r"[\s_]+", "_", text).strip("_")
    return text.lower() or "unknown"


def ensure_dir(path: Union[str, Path]) -> Path:
    p = Path(path)
    p.mkdir(parents=True, exist_ok=True)
    return p


def read_jsonl(path: Union[str, Path]) -> Iterator[Dict[str, Any]]:
    with open(path, "r", encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if line:
                yield json.loads(line)


def write_json(path: Union[str, Path], data: Any) -> None:
    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


def write_jsonl(path: Union[str, Path], rows: Iterable[Dict[str, Any]]) -> None:
    with open(path, "w", encoding="utf-8") as f:
        for row in rows:
            f.write(json.dumps(row, ensure_ascii=False) + "\n")


def iter_docx_blocks(doc: docx.Document):
    """Yield paragraphs and tables in document order."""
    def walk(element):
        for child in element:
            if child.tag.endswith("}p"):
                yield docx.text.paragraph.Paragraph(child, doc)
            elif child.tag.endswith("}tbl"):
                yield docx.table.Table(child, doc)
            else:
                yield from walk(child)
    yield from walk(doc.element.body)


def get_docx_texts(docx_path: Union[str, Path], limit: Optional[int] = None) -> List[str]:
    doc = docx.Document(str(docx_path))
    texts = [collapse_ws(p.text) for p in doc.paragraphs if collapse_ws(p.text)]
    return texts[:limit] if limit else texts


def get_docx_block_texts(docx_path: Union[str, Path], limit: Optional[int] = None) -> List[str]:
    doc = docx.Document(str(docx_path))
    texts: List[str] = []
    for block in iter_docx_blocks(doc):
        if isinstance(block, docx.text.paragraph.Paragraph):
            text = collapse_ws(block.text)
            if text:
                texts.append(text)
        elif isinstance(block, docx.table.Table):
            for row in block.rows:
                for cell in row.cells:
                    text = collapse_ws(cell.text)
                    if text:
                        texts.append(text)
        if limit and len(texts) >= limit:
            break
    return texts[:limit] if limit else texts


def _docx_table_cell_text(cell: docx.table._Cell) -> str:
    paragraphs = [collapse_ws(p.text) for p in cell.paragraphs if collapse_ws(p.text)]
    return "\n".join(paragraphs) if paragraphs else collapse_ws(cell.text)


def _docx_tables_to_html(docx_path: Union[str, Path]) -> List[str]:
    try:
        document = docx.Document(str(docx_path))
    except Exception:
        return []

    tables: List[str] = []
    for table in document.tables:
        row_html = []
        for row in table.rows:
            cells = []
            for cell in row.cells:
                text = html_lib.escape(_docx_table_cell_text(cell))
                cells.append(f"<td>{text}</td>")
            row_html.append("<tr>" + "".join(cells) + "</tr>")
        tables.append(
            "<table border=\"1\" style=\"border-collapse: collapse; width: 100%;\">"
            + "".join(row_html)
            + "</table>"
        )
    return tables


def get_mammoth_html_tables(docx_path: Union[str, Path]) -> List[str]:
    if mammoth is None or BeautifulSoup is None:
        return _docx_tables_to_html(docx_path)
    try:
        with open(docx_path, "rb") as f:
            result = mammoth.convert_to_html(f)
    except Exception:
        return _docx_tables_to_html(docx_path)

    soup = BeautifulSoup(result.value, "html.parser")
    tables = soup.find_all("table")
    out = []
    for tbl in tables:
        tbl["border"] = "1"
        tbl["style"] = "border-collapse: collapse; width: 100%;"
        for cell in tbl.find_all(["td", "th"]):
            paragraphs = [collapse_ws(p.get_text(" ", strip=True)) for p in cell.find_all("p")]
            if paragraphs:
                cell.clear()
                cell.append("\n".join([p for p in paragraphs if p]))
            else:
                txt = collapse_ws(cell.get_text(" ", strip=True))
                cell.clear()
                cell.append(txt)
        out.append(str(tbl))
    return out


def normalize_html_table(html: str) -> List[List[str]]:
    if BeautifulSoup is None:
        return []
    soup = BeautifulSoup(html or "", "html.parser")
    table = soup.find("table")
    if table is None:
        return []

    grid: List[List[Optional[str]]] = []
    carry: Dict[Tuple[int, int], str] = {}
    max_cols = 0

    for r_idx, tr in enumerate(table.find_all("tr")):
        grid.append([])
        c_idx = 0
        while (r_idx, c_idx) in carry:
            grid[r_idx].append(carry[(r_idx, c_idx)])
            c_idx += 1

        for cell in tr.find_all(["td", "th"]):
            while (r_idx, c_idx) in carry:
                grid[r_idx].append(carry[(r_idx, c_idx)])
                c_idx += 1

            txt = collapse_ws(cell.get_text(" ", strip=True))
            colspan = int(cell.get("colspan", 1) or 1)
            rowspan = int(cell.get("rowspan", 1) or 1)

            for dc in range(colspan):
                grid[r_idx].append(txt)
                if rowspan > 1:
                    for dr in range(1, rowspan):
                        carry[(r_idx + dr, c_idx + dc)] = txt
            c_idx += colspan

        while (r_idx, c_idx) in carry:
            grid[r_idx].append(carry[(r_idx, c_idx)])
            c_idx += 1

        max_cols = max(max_cols, len(grid[r_idx]))

    rows = []
    for row in grid:
        filled = [(x or "") for x in row]
        filled.extend([""] * (max_cols - len(filled)))
        rows.append(filled)
    return rows


def normalize_document_number(text: str) -> Optional[str]:
    if not text:
        return None
    raw = collapse_ws(text).replace("\u00a0", " ")
    m = re.search(
        r"(?P<num>\d+)\s*/\s*(?P<year>\d{4})\s*/\s*(?P<kind>[A-ZĐ]+)"
        r"(?:\s*[-–—]?\s*(?P<agency>[A-ZĐ0-9]+(?:\s*[-–—]\s*[A-ZĐ0-9]+)*))?"
        r"(?=\s*(?:ngày|,|;|\.|\)|$))",
        raw, re.I | re.U
    )
    if not m:
        return None
    num = m.group("num")
    year = m.group("year")
    kind = (m.group("kind") or "").upper().replace("Đ", "D")
    agency = (m.group("agency") or "").upper().replace("Đ", "D")
    agency = re.sub(r"[^A-Z0-9]+", "", agency)
    return f"{num}/{year}/{kind}-{agency}" if agency else f"{num}/{year}/{kind}"


def extract_issued_with(text: str) -> Optional[str]:
    return normalize_document_number(text)


def make_text_for_embedding(*, path_text: str, content: str, title: str = "", extra: str = "") -> str:
    parts = []
    if title:
        parts.append(f"Tiêu đề: {title}")
    if path_text:
        parts.append(f"Đường dẫn: {path_text}")
    if content:
        parts.append(f"Nội dung: {content}")
    if extra:
        parts.append(extra)
    return "\n".join(parts)


def maybe_extract_label_from_filename(path: Union[str, Path]) -> Optional[str]:
    name = Path(path).stem
    qcvn_code = normalize_qcvn_code(name)
    if qcvn_code:
        return qcvn_code
    m = re.search(r"(Phụ\s*lục|Phu\s*luc)\s+([IVXLCDM]+[A-Z]?|\d+[A-Z]?|[A-Z]+)", name, re.I | re.U)
    if m:
        token = m.group(2).strip().strip("._- ")
        return f"Phụ lục {token.upper()}"
    m = re.search(r"(Mẫu|Mau)\s+(?:số\s+|so\s+)?([0-9A-Za-zĐđ_.-]+)", name, re.I | re.U)
    if m:
        token = m.group(2).strip().strip("._- ")
        return f"Mẫu số {token}"
    name_ascii = strip_vietnamese_accents(name, keep_dd=False)
    m = re.match(r"^\s*(?P<token>\d{1,3}[A-Za-z]?)\b", name_ascii, re.I | re.U)
    if m:
        return f"Mẫu số {m.group('token').upper()}"
    return None


def normalize_qcvn_code(text: str) -> Optional[str]:
    raw = collapse_ws(text or "").replace("_", " ")
    m = re.search(
        r"\bQCVN\s*(?P<num>\d+[A-Z]?)\s*(?::|\s)\s*(?P<year>\d{4})\s*(?:/|\s)\s*(?P<agency>[A-ZĐ]+)\b",
        raw,
        re.I | re.U,
    )
    if not m:
        return None
    agency = m.group("agency").upper().replace("Đ", "D")
    return f"QCVN {m.group('num')}:{m.group('year')}/{agency}"


def extract_qcvn_header(texts: List[str], source_file: Union[str, Path]) -> Optional[Dict[str, Any]]:
    haystack = [Path(source_file).stem] + list(texts[:80])
    label = None
    for t in haystack:
        label = normalize_qcvn_code(t)
        if label:
            break
    if not label:
        return None

    title = None
    for idx, text in enumerate(texts[:80]):
        title = _qcvn_title_from_line_window(texts, idx)
        if title:
            break

    return {
        "label": label,
        "title": title or Path(source_file).stem,
        "issued_with": extract_issued_with("\n".join(texts[:20])),
    }


def _qcvn_title_from_line_window(texts: List[str], start_idx: int) -> Optional[str]:
    first = _qcvn_title_segment(texts[start_idx])
    if not first:
        return None

    parts = [first]
    for value in texts[start_idx + 1:start_idx + 8]:
        text = collapse_ws(value)
        if not text or _is_qcvn_title_stop_line(text):
            break
        if _looks_like_vietnamese_title_continuation(text):
            parts.append(text)
            continue
        break
    return _clean_qcvn_title(" ".join(parts))


def _qcvn_title_segment(text: str) -> Optional[str]:
    raw = collapse_ws(text)
    ascii_raw = strip_vietnamese_accents(raw, keep_dd=False).lower()
    marker = "quy chuan ky thuat quoc gia"
    marker_idx = ascii_raw.find(marker)
    if marker_idx < 0:
        return None

    prefix_ascii = strip_vietnamese_accents(raw[:marker_idx], keep_dd=False).lower()
    if "danh muc" in prefix_ascii:
        return None

    segment = raw[marker_idx:]
    segment = re.split(
        r"\bNational\s+technical\b|\bNational\s+Technical\b|\bH[àa]\s*Nội\b|\bHa\s+Noi\b|\bLời\s+nói\s+đầu\b|\bMục\s+lục\b",
        segment,
        maxsplit=1,
        flags=re.I | re.U,
    )[0]
    return collapse_ws(segment)


def _is_qcvn_title_stop_line(text: str) -> bool:
    ascii_text = strip_vietnamese_accents(text, keep_dd=False).lower()
    if re.match(r"^qcvn\b", ascii_text):
        return True
    if re.match(r"^(national technical|ha noi|loi noi dau|muc luc|trang)\b", ascii_text):
        return True
    if re.match(r"^(\d+(?:\.\d+)*|[ivxlcdm]+)\.\s+", text, re.I | re.U):
        return True
    if re.fullmatch(r"[_\-\s\.]+", text):
        return True
    if "cong hoa xa hoi" in ascii_text:
        return True
    return False


def _looks_like_vietnamese_title_continuation(text: str) -> bool:
    letters = [ch for ch in text if ch.isalpha()]
    if not letters:
        return False
    ascii_text = strip_vietnamese_accents(text, keep_dd=False).lower()
    if re.match(r"^(national technical|ha noi|loi noi dau|muc luc|trang)\b", ascii_text):
        return False
    upper_ratio = sum(1 for ch in letters if ch.isupper()) / len(letters)
    return upper_ratio >= 0.65


def _clean_qcvn_title(text: str) -> str:
    title = collapse_ws(text)
    title = re.sub(r"\bQCVN\s*\d+[A-Z]?\s*:\s*\d{4}\s*/\s*[A-ZĐ]+\b", "", title, flags=re.I | re.U)
    title = re.sub(r"\bH[àa]\s*Nội\s*[-–]\s*\d{4}\b", "", title, flags=re.I | re.U)
    title = re.sub(r"\bHa\s+Noi\s*[-–]\s*\d{4}\b", "", title, flags=re.I | re.U)
    title = re.sub(r"(QUY CHUẨN KỸ THUẬT QUỐC GIA VỀ)\s+VỀ\b", r"\1", title, flags=re.I | re.U)
    title = re.sub(r"\s+([.,;:])", r"\1", title)
    return collapse_ws(title).strip(" .")


def extract_attachment_header(texts: List[str], source_file: Union[str, Path]) -> Dict[str, Any]:
    qcvn_header = extract_qcvn_header(texts, source_file)
    if qcvn_header:
        return qcvn_header

    label = None
    title = None
    issued_with = None

    for t in texts[:8]:
        m = re.match(r"^(Phụ\s*lục|Phu\s*luc)\s+([IVXLCDM]+[A-Z]?|\d+[A-Z]?|[A-Z]+)\b", t, re.I | re.U)
        if m:
            token = m.group(2).strip().strip("._- ")
            label = f"Phụ lục {token.upper()}"
            break
        if re.match(r"^(Phụ\s*lục|Phu\s*luc)\b", t, re.I | re.U):
            label = "Phụ lục"
            break
    if not label:
        label = maybe_extract_label_from_filename(source_file)

    joined_head = "\n".join(texts[:12])
    issued_with = extract_issued_with(joined_head)

    def is_separator_line(value: str) -> bool:
        return bool(re.fullmatch(r"[_\-\s\.]+", value or ""))

    def is_structured_heading(value: str) -> bool:
        return bool(re.match(r"^([A-ZĐ]|[IVXLCDM]+|\d+(?:\.\d+)*)\.\s+", value or "", re.I | re.U))

    def is_boilerplate_title_line(value: str) -> bool:
        low = (value or "").lower()
        up = (value or "").upper()
        if not value or is_separator_line(value):
            return True
        if label and low.startswith(label.lower()):
            return True
        if any(x in up for x in ["BAN HÀNH KÈM", "KÈM THEO", "THÔNG TƯ SỐ", "NGHỊ ĐỊNH SỐ", "QUYẾT ĐỊNH SỐ"]):
            return True
        if "CỦA BỘ TRƯỞNG" in up or "CỦA CHÍNH PHỦ" in up:
            return True
        if re.match(r"^ngày\s+\d{1,2}\s+tháng\s+\d{1,2}\s+năm\s+\d{4}", low, re.U):
            return True
        if low.startswith("của "):
            return True
        if "CỘNG HÒA" in up or "ĐỘC LẬP" in up:
            return True
        return False

    def is_strong_title_line(value: str) -> bool:
        letters = [ch for ch in value if ch.isalpha()]
        if len(value) < 8 or not letters:
            return False
        upper_ratio = sum(1 for ch in letters if ch.isupper()) / len(letters)
        return upper_ratio >= 0.7

    def append_title_part(parts: List[str], value: str) -> None:
        current = strip_vietnamese_accents(value, keep_dd=False).lower()
        for idx, existing in enumerate(parts):
            old = strip_vietnamese_accents(existing, keep_dd=False).lower()
            if current == old or current in old:
                return
            if old in current:
                parts[idx] = value
                return
        parts.append(value)

    def maybe_title_from_form_heading(value: str) -> Optional[str]:
        nonlocal label
        m = re.match(
            r"^(Mẫu|Mau)\s+(?:số\s+|so\s+)?(?P<no>[0-9A-Za-zĐđ_.-]+)\.?\s*(?P<title>.*)",
            value or "",
            re.I | re.U,
        )
        if not m:
            return None
        form_no = m.group("no").strip().strip("._- ")
        if not re.search(r"\d", form_no):
            return None
        nonlocal_label = f"Mẫu số {form_no}"
        if not label:
            label = nonlocal_label
        tail = collapse_ws(m.group("title"))
        if not tail or is_boilerplate_title_line(tail):
            return None
        return tail

    for t in texts[:8]:
        form_title = maybe_title_from_form_heading(t)
        if form_title:
            title = form_title
            break

    # Title heuristic: collect title lines, including multi-line titles separated
    # by issued-with boilerplate. Stop before the first structured body heading.
    if not title:
        title_parts = []
        for t in texts[:20]:
            if is_boilerplate_title_line(t):
                continue
            if is_structured_heading(t):
                if title_parts:
                    break
                continue
            if is_strong_title_line(t):
                append_title_part(title_parts, t)

        if title_parts:
            title = collapse_ws(" ".join(title_parts))

    if not title:
        plain_parts = []
        for t in texts[:20]:
            if is_boilerplate_title_line(t):
                continue
            low_ascii = strip_vietnamese_accents(t, keep_dd=False).lower()
            if is_structured_heading(t) or re.match(r"^bang\s+\d+", low_ascii):
                if plain_parts:
                    break
                continue
            if len(t) >= 8:
                plain_parts.append(t)
            if len(plain_parts) >= 3:
                break
        if plain_parts:
            title = collapse_ws(" ".join(plain_parts))

    if not title:
        # fallback to filename tail
        title = Path(source_file).stem

    return {
        "label": label,
        "title": collapse_ws(title),
        "issued_with": issued_with,
    }


def extract_ref_mentions_light(text: str) -> List[Dict[str, Any]]:
    """Lightweight ref mention extraction for attachments."""
    mentions: List[Dict[str, Any]] = []
    patterns = [
        ("appendix", re.compile(r"\b[Pp]hụ\s+[Ll]ục\s+(?P<label>[IVXLCDM]+[A-Z]?|\d+[A-Z]?|[A-Z]+)\b", re.U)),
        ("form", re.compile(r"\b[Mm]ẫu\s+(?:s\S*\s+)?(?P<label>[0-9A-Za-zĐđ_.\-/]+)", re.U)),
        ("article", re.compile(r"\b[Đđ]iều\s+(?P<label>\d+[a-zA-Z]?)\b", re.U)),
        ("legal_document", re.compile(r"\b(?:Thông\s+tư|Nghị\s+định|Luật|Quyết\s+định|Nghị\s+quyết)\s+số\s+(?P<label>[0-9]+/[0-9A-ZĐđ\-_/\.]+)", re.I | re.U)),
    ]
    seen = set()
    for typ, pat in patterns:
        for m in pat.finditer(text or ""):
            raw = collapse_ws(m.group(0))
            label = collapse_ws(m.group("label")).strip(".,;:)")
            if typ == "form" and not re.search(r"\d", label):
                continue
            if typ == "legal_document":
                label = label.upper()
            key = (typ, raw, m.start(), m.end())
            if key in seen:
                continue
            seen.add(key)
            mentions.append({
                "mention_type": typ,
                "label": label,
                "raw": raw,
                "span": [m.start(), m.end()],
                "needs_resolution": True,
                "resolution": {"status": "unresolved", "target_id": None, "target_type": None, "resolver": None, "confidence": None},
            })
    return mentions


def is_probable_field_line(text: str) -> bool:
    if not text:
        return False
    if "……" in text or "..." in text or "…" in text:
        return True
    # Label ending colon, not too long.
    if re.match(r"^[\-–•]?\s*[A-ZÀ-ỸĐa-zà-ỹđ0-9 ,/()]+:\s*.*$", text) and len(text) <= 140:
        return True
    return False


def field_label_from_line(text: str) -> Optional[str]:
    t = re.sub(r"^[\-–•]\s*", "", collapse_ws(text))
    m = re.match(r"(?P<label>[^:：]{2,80})[:：]", t)
    if m:
        return collapse_ws(m.group("label"))
    if "…" in t or "……" in t or "..." in t:
        before = re.split(r"…+|\.{3,}", t, maxsplit=1)[0]
        before = collapse_ws(before).rstrip(":：")
        if 2 <= len(before) <= 80:
            return before
    return None
