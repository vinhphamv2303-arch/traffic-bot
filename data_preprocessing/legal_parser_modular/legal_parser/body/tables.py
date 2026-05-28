from __future__ import annotations
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Union
import html as html_lib
import docx
import mammoth
from bs4 import BeautifulSoup
from ..common.utils import collapse_ws

def _docx_table_cell_text(cell: docx.table._Cell) -> str:
    paragraphs = [collapse_ws(p.text) for p in cell.paragraphs if collapse_ws(p.text)]
    return "\n".join(paragraphs) if paragraphs else collapse_ws(cell.text)

def _docx_tables_to_html(docx_path: Union[str, Path]) -> List[str]:
    try:
        document = docx.Document(str(docx_path))
    except Exception:
        return []

    html_tables: List[str] = []
    for table in document.tables:
        row_html = []
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cells.append(f"<td>{html_lib.escape(_docx_table_cell_text(cell))}</td>")
            row_html.append("<tr>" + "".join(cells) + "</tr>")
        html_tables.append(
            "<table border=\"1\" style=\"border-collapse: collapse; width: 100%;\">"
            + "".join(row_html)
            + "</table>"
        )
    return html_tables

def get_mammoth_html_tables(docx_path: Union[str, Path]) -> List[str]:
    try:
        with open(docx_path, "rb") as docx_file:
            result = mammoth.convert_to_html(docx_file)
    except Exception:
        return _docx_tables_to_html(docx_path)

    soup = BeautifulSoup(result.value, "html.parser")
    html_tables: List[str] = []
    for tbl in soup.find_all("table"):
        tbl["border"] = "1"; tbl["style"] = "border-collapse: collapse; width: 100%;"
        for cell in tbl.find_all(["td", "th"]):
            ps = [collapse_ws(p.get_text(" ", strip=True)) for p in cell.find_all("p")]
            cell.clear(); cell.append("\n".join([p for p in ps if p]) if ps else collapse_ws(cell.get_text(" ", strip=True)))
        html_tables.append(str(tbl))
    return html_tables

def normalize_html_table(html: str) -> List[List[str]]:
    soup = BeautifulSoup(html or "", "html.parser")
    table = soup.find("table")
    if table is None: return []
    grid: List[List[Optional[str]]] = []
    carry: Dict[Tuple[int, int], str] = {}
    max_cols = 0
    for r_idx, tr in enumerate(table.find_all("tr")):
        grid.append([]); c_idx = 0
        while (r_idx, c_idx) in carry:
            grid[r_idx].append(carry[(r_idx, c_idx)]); c_idx += 1
        for cell in tr.find_all(["td", "th"]):
            while (r_idx, c_idx) in carry:
                grid[r_idx].append(carry[(r_idx, c_idx)]); c_idx += 1
            txt = collapse_ws(cell.get_text(" ", strip=True))
            colspan = int(cell.get("colspan", 1) or 1); rowspan = int(cell.get("rowspan", 1) or 1)
            for dc in range(colspan):
                grid[r_idx].append(txt)
                if rowspan > 1:
                    for dr in range(1, rowspan): carry[(r_idx + dr, c_idx + dc)] = txt
            c_idx += colspan
        while (r_idx, c_idx) in carry:
            grid[r_idx].append(carry[(r_idx, c_idx)]); c_idx += 1
        max_cols = max(max_cols, len(grid[r_idx]))
    out=[]
    for row in grid:
        filled=[x or "" for x in row]; filled.extend([""]*(max_cols-len(filled))); out.append(filled)
    return out
