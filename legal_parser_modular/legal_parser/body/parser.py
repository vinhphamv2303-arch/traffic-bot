from __future__ import annotations

import re
from pathlib import Path
from typing import Any, Dict, Iterable, Iterator, List, Optional, Union

import docx

from .docx_iter import iter_blocks_recursive
from .mentions import extract_amendment_mentions, extract_ref_mentions, flags_for_text
from ..common.models import ParseResult, ParserConfig
from .tables import get_mammoth_html_tables, normalize_html_table
from ..common.utils import clean_filename, collapse_ws, ensure_dir, md5_text, normalize_id, normalize_document_number, \
    write_json, write_jsonl


class LegalBodyParser:
    pattern_phan = re.compile(r"^Phần\s+(thứ\s+[a-zđà-ỹA-ZÀ-Ỹ]+|[IVXLCDM\d]+)[.:\-]?\s*(.*)", re.IGNORECASE)
    pattern_chuong = re.compile(r"^Chương\s+([IVXLCDM\d]+)[.:\-]?\s*(.*)", re.IGNORECASE)
    pattern_muc = re.compile(r"^Mục\s+(\d+|[IVXLCDM]+)[.:\-]?\s*(.*)", re.IGNORECASE)
    pattern_tieu_muc = re.compile(r"^Tiểu\s*mục\s+(\d+|[IVXLCDM]+)[.:\-]?\s*(.*)", re.IGNORECASE)
    pattern_dieu = re.compile(r"^Điều\s+(\d+[a-zA-Z]*)[.:]?\s*(.*)", re.IGNORECASE)
    pattern_khoan = re.compile(r"^(\d+)\.\s*(.*)")
    pattern_diem = re.compile(r"^([a-zđ])\)\s*(.*)", re.IGNORECASE)
    doc_type_pattern = re.compile(r"^(LUẬT|BỘ\s+LUẬT|NGHỊ\s+ĐỊNH|THÔNG\s+TƯ|QUYẾT\s+ĐỊNH|NGHỊ\s+QUYẾT)$", re.IGNORECASE)
    issue_date_pattern = re.compile(r"\bngày\s+\d{1,2}\s+tháng\s+\d{1,2}\s+năm\s+\d{4}\b", re.IGNORECASE)

    def __init__(self, config: Optional[ParserConfig] = None):
        self.config = config or ParserConfig()

    @classmethod
    def extract_issue_date_text(cls, text: str) -> str:
        m = cls.issue_date_pattern.search(text or "")
        return collapse_ws(m.group(0)) if m else ""

    def extract_metadata(self, doc: docx.Document, docx_path: Union[str, Path]) -> Dict[str, Any]:
        metadata: Dict[str, Any] = {
            "document_number": "UNKNOWN",
            "document_id": "unknown",
            "document_type": "",
            "issue_date": "",
            "issuing_agency": "",
            "national_motto": "",
            "document_title": "",
            "preamble": [],
            "source_file": str(docx_path),
        }

        # Regex bắt số hiệu văn bản nhưng KHÔNG làm mất hậu tố như TT-BCA, TT-BXD, NĐ-CP...
        doc_number_pattern = re.compile(
            r"\bSố\s*:\s*"
            r"([0-9]+\s*/\s*[0-9]{4}\s*/\s*[A-ZĐ]+(?:\s*[-–—]?\s*[A-ZĐ0-9]+)*)",
            re.IGNORECASE | re.UNICODE,
        )

        # 1. Ưu tiên đọc header table
        if len(doc.tables) > 0:
            for row in doc.tables[0].rows:
                for cell in row.cells:
                    # QUAN TRỌNG:
                    # Không được thay '-' thành khoảng trắng ở raw_cell_text,
                    # nếu không 12/2025/TT-BCA sẽ bị thành 12/2025/TT BCA
                    # rồi regex cũ chỉ bắt được 12/2025/TT.
                    raw_cell_text = collapse_ws((cell.text or "").replace("\u00a0", " "))
                    text = collapse_ws(re.sub(r"[_]+", " ", raw_cell_text))

                    if not text:
                        continue

                    if metadata["document_number"] == "UNKNOWN":
                        m = doc_number_pattern.search(raw_cell_text)
                        if m:
                            metadata["document_number"] = normalize_document_number(m.group(1))

                    if not metadata["issue_date"] and "ngày" in text.lower() and "tháng" in text.lower():
                        metadata["issue_date"] = self.extract_issue_date_text(text)

                    upper = text.upper()
                    if (
                            not metadata["issuing_agency"]
                            and (
                            "BỘ" in upper
                            or "CHÍNH PHỦ" in upper
                            or "QUỐC HỘI" in upper
                            or "ỦY BAN" in upper
                    )
                    ):
                        metadata["issuing_agency"] = collapse_ws((cell.text or "").split("\n")[0])

                    if not metadata["national_motto"] and "CỘNG HÒA" in upper:
                        metadata["national_motto"] = collapse_ws((cell.text or "").split("\n")[0])

        # 2. Fallback: đọc 60 paragraph đầu nếu header table không bắt được
        for text in [collapse_ws(p.text) for p in doc.paragraphs[:60] if collapse_ws(p.text)]:
            if metadata["document_number"] == "UNKNOWN":
                m = doc_number_pattern.search(text)
                if m:
                    metadata["document_number"] = normalize_document_number(m.group(1))

            if not metadata["issue_date"] and "ngày" in text.lower() and "tháng" in text.lower():
                metadata["issue_date"] = self.extract_issue_date_text(text)

            if not metadata["issuing_agency"]:
                upper = text.upper()
                if upper in {"BỘ CÔNG AN", "CHÍNH PHỦ", "QUỐC HỘI"} or upper.startswith("BỘ "):
                    metadata["issuing_agency"] = text

            if not metadata["national_motto"] and "CỘNG HÒA" in text.upper():
                metadata["national_motto"] = text

        base = metadata["document_number"] if metadata["document_number"] != "UNKNOWN" else Path(docx_path).stem
        metadata["document_id"] = normalize_id(base)

        return metadata

    @staticmethod
    def node_label(node_type: str, no: str, content: str) -> str:
        prefixes = {"phan": "Phần", "chuong": "Chương", "muc": "Mục", "tieu_muc": "Tiểu mục", "dieu": "Điều",
                    "khoan": "Khoản", "diem": "Điểm", "text": "Đoạn", "table": "Bảng"}
        prefix = prefixes.get(node_type, node_type)
        if node_type in {"text", "table"}:
            return f"{prefix} {no}"
        return collapse_ws(f"{prefix} {no}. {content}" if content else f"{prefix} {no}")

    def make_node(self, *, document_id: str, node_type: str, no: str, content: str, parent: Optional[Dict[str, Any]],
                  path_stack: List[Dict[str, str]], node_id: Optional[str] = None,
                  extra: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
        norm_no = normalize_id(no)
        node_id = node_id or (
            f"{parent['id']}.{node_type}_{norm_no}" if parent else f"{document_id}.{node_type}_{norm_no}")
        label = self.node_label(node_type, no, content)
        path = [x["id"] for x in path_stack] + [node_id]
        path_text = " > ".join([*[x["label"] for x in path_stack], label])
        ref_mentions = extract_ref_mentions(content)
        amendment_mentions = extract_amendment_mentions(content)
        node = {
            "id": node_id, "document_id": document_id, "type": node_type, "no": norm_no, "raw_no": no,
            "content": content, "parent_id": parent["id"] if parent else None, "path": path, "path_text": path_text,
            "label": label,
            "flags": {"has_ref_or_amend": bool(ref_mentions or amendment_mentions),
                      "has_reference_mention": bool(ref_mentions), "has_amendment_mention": bool(amendment_mentions),
                      "needs_reference_resolution": bool(ref_mentions),
                      "needs_amendment_processing": bool(amendment_mentions)},
            "ref_mentions": ref_mentions, "amendment_mentions": amendment_mentions, "children": [],
        }
        if extra:
            node.update(extra)
        return node

    @staticmethod
    def attach_node(node: Dict[str, Any], parent: Optional[Dict[str, Any]], root: List[Dict[str, Any]]) -> None:
        (parent.setdefault("children", []) if parent else root).append(node)

    @staticmethod
    def current_parent(state: Dict[str, Optional[Dict[str, Any]]]) -> Optional[Dict[str, Any]]:
        for key in ["diem", "khoan", "dieu", "tieu_muc", "muc", "chuong", "phan"]:
            if state.get(key):
                return state[key]
        return None

    @staticmethod
    def path_stack_from_state(metadata: Dict[str, Any], state: Dict[str, Optional[Dict[str, Any]]],
                              stop_before: Optional[str] = None) -> List[Dict[str, str]]:
        stack = [{"id": metadata["document_id"], "label": metadata.get("document_number") or metadata["document_id"]}]
        for key in ["phan", "chuong", "muc", "tieu_muc", "dieu", "khoan", "diem"]:
            if key == stop_before:
                break
            node = state.get(key)
            if node:
                stack.append({"id": node["id"], "label": node.get("label", node["id"])})
        return stack

    def parse_file(self, docx_path: Union[str, Path]) -> ParseResult:
        docx_path = Path(docx_path)
        doc = docx.Document(str(docx_path))
        html_tables = get_mammoth_html_tables(docx_path)
        table_idx = 0
        metadata = self.extract_metadata(doc, docx_path)
        document_id = metadata["document_id"]
        out_name = clean_filename(
            metadata["document_number"] if metadata["document_number"] != "UNKNOWN" else docx_path.stem)
        doc_dir = ensure_dir(self.config.output_base_dir / out_name)
        content_tree: List[Dict[str, Any]] = []
        footer_data: List[Dict[str, Any]] = []
        tables_flat: List[Dict[str, Any]] = []
        used_node_ids: set[str] = set()
        state: Dict[str, Optional[Dict[str, Any]]] = {k: None for k in
                                                      ["phan", "chuong", "muc", "tieu_muc", "dieu", "khoan", "diem"]}
        ps: Dict[str, Any] = {"in_quote_block": False, "pending_title_node": None, "quote_buffer": [],
                              "is_preamble": True, "is_footer": False, "collecting_title": False}

        def reset_below(level: str) -> None:
            order = ["phan", "chuong", "muc", "tieu_muc", "dieu", "khoan", "diem"]
            for k in order[order.index(level) + 1:]: state[k] = None

        def unique_node_id(base_id: str) -> str:
            node_id = base_id
            suffix = 2
            while node_id in used_node_ids:
                node_id = f"{base_id}_{suffix}"
                suffix += 1
            used_node_ids.add(node_id)
            return node_id

        def child_node_id(parent: Optional[Dict[str, Any]], node_type: str, no: str) -> str:
            parent_id = parent["id"] if parent else document_id
            return unique_node_id(f"{parent_id}.{node_type}_{normalize_id(no)}")

        def append_text_node(text: str) -> None:
            parent = self.current_parent(state);
            parent_id = parent["id"] if parent else document_id
            siblings = parent.get("children", []) if parent else content_tree
            idx = 1 + sum(1 for c in siblings if c.get("type") == "text")
            node = self.make_node(document_id=document_id, node_type="text", no=str(idx), content=text, parent=parent,
                                  path_stack=self.path_stack_from_state(metadata, state),
                                  node_id=unique_node_id(f"{parent_id}.text_{idx}"))
            self.attach_node(node, parent, content_tree)

        def flush_quote() -> None:
            if ps["quote_buffer"]:
                append_text_node("\n".join(ps["quote_buffer"]));
                ps["quote_buffer"] = [];
                ps["in_quote_block"] = False

        def handle_pending(text: str, is_structure: bool) -> bool:
            pending = ps.get("pending_title_node")
            if pending is not None and not is_structure:
                pending["content"] = collapse_ws((pending.get("content") or "") + " " + text)
                pending["label"] = self.node_label(pending["type"], pending.get("raw_no", pending["no"]),
                                                   pending["content"])
                pending["ref_mentions"] = extract_ref_mentions(pending["content"])
                pending["amendment_mentions"] = extract_amendment_mentions(pending["content"])
                pending["flags"] = flags_for_text(pending["content"])
                ps["pending_title_node"] = None
                return True
            return False

        def process_text(text: str) -> None:
            text = collapse_ws(text)
            open_q, close_q, std_q = text.count("“"), text.count("”"), text.count('"')
            was = ps["in_quote_block"]
            if open_q > close_q or (std_q % 2 and not was):
                ps["in_quote_block"] = True
            elif close_q > open_q or (std_q % 2 and was):
                ps["in_quote_block"] = False
            elif text.endswith("”") or text.endswith('"'):
                ps["in_quote_block"] = False
            if was or text.startswith("“") or text.startswith('"'):
                ps["pending_title_node"] = None;
                ps["quote_buffer"].append(text)
                if not ps["in_quote_block"]: flush_quote()
                return
            matches = {"phan": self.pattern_phan.match(text), "chuong": self.pattern_chuong.match(text),
                       "muc": self.pattern_muc.match(text), "tieu_muc": self.pattern_tieu_muc.match(text),
                       "dieu": self.pattern_dieu.match(text), "khoan": self.pattern_khoan.match(text),
                       "diem": self.pattern_diem.match(text)}
            is_structure = any(matches.values())
            if handle_pending(text, is_structure): return
            if is_structure: ps["pending_title_node"] = None
            if matches["phan"]:
                so, content = matches["phan"].group(1), collapse_ws(matches["phan"].group(2));
                parent = None
                node = self.make_node(document_id=document_id, node_type="phan", no=so, content=content, parent=parent,
                                      path_stack=self.path_stack_from_state(metadata, state, "phan"),
                                      node_id=unique_node_id(f"{document_id}.phan_{normalize_id(so)}"))
                self.attach_node(node, parent, content_tree);
                state["phan"] = node;
                reset_below("phan")
                if not content: ps["pending_title_node"] = node
                return
            if matches["chuong"]:
                so, content = matches["chuong"].group(1), collapse_ws(matches["chuong"].group(2));
                parent = state["phan"]
                node = self.make_node(document_id=document_id, node_type="chuong", no=so, content=content,
                                      parent=parent, path_stack=self.path_stack_from_state(metadata, state, "chuong"),
                                      node_id=unique_node_id(f"{document_id}.chuong_{normalize_id(so)}"))
                self.attach_node(node, parent, content_tree);
                state["chuong"] = node;
                reset_below("chuong")
                if not content: ps["pending_title_node"] = node
                return
            if matches["muc"]:
                so, content = matches["muc"].group(1), collapse_ws(matches["muc"].group(2));
                parent = state["chuong"] or state["phan"]
                node = self.make_node(document_id=document_id, node_type="muc", no=so, content=content, parent=parent,
                                      path_stack=self.path_stack_from_state(metadata, state, "muc"),
                                      node_id=child_node_id(parent, "muc", so))
                self.attach_node(node, parent, content_tree);
                state["muc"] = node;
                reset_below("muc")
                if not content: ps["pending_title_node"] = node
                return
            if matches["tieu_muc"]:
                so, content = matches["tieu_muc"].group(1), collapse_ws(matches["tieu_muc"].group(2));
                parent = state["muc"] or state["chuong"] or state["phan"]
                node = self.make_node(document_id=document_id, node_type="tieu_muc", no=so, content=content,
                                      parent=parent, path_stack=self.path_stack_from_state(metadata, state, "tieu_muc"),
                                      node_id=child_node_id(parent, "tieu_muc", so))
                self.attach_node(node, parent, content_tree);
                state["tieu_muc"] = node;
                reset_below("tieu_muc")
                if not content: ps["pending_title_node"] = node
                return
            if matches["dieu"]:
                so, content = matches["dieu"].group(1), collapse_ws(matches["dieu"].group(2));
                parent = state["tieu_muc"] or state["muc"] or state["chuong"] or state["phan"]
                node = self.make_node(document_id=document_id, node_type="dieu", no=so, content=content, parent=parent,
                                      path_stack=self.path_stack_from_state(metadata, state, "dieu"),
                                      node_id=unique_node_id(f"{document_id}.dieu_{normalize_id(so)}"))
                self.attach_node(node, parent, content_tree);
                state["dieu"] = node;
                reset_below("dieu")
                if not content: ps["pending_title_node"] = node
                return
            if matches["khoan"] and state["dieu"]:
                so, content = matches["khoan"].group(1), collapse_ws(matches["khoan"].group(2));
                parent = state["dieu"]
                node = self.make_node(document_id=document_id, node_type="khoan", no=so, content=content, parent=parent,
                                      path_stack=self.path_stack_from_state(metadata, state, "khoan"),
                                      node_id=child_node_id(parent, "khoan", so))
                self.attach_node(node, parent, content_tree);
                state["khoan"] = node;
                reset_below("khoan");
                return
            if matches["diem"] and (state["khoan"] or state["dieu"]):
                so, content = matches["diem"].group(1), collapse_ws(matches["diem"].group(2));
                parent = state["khoan"] or state["dieu"]
                node = self.make_node(document_id=document_id, node_type="diem", no=so, content=content, parent=parent,
                                      path_stack=self.path_stack_from_state(metadata, state, "diem"),
                                      node_id=child_node_id(parent, "diem", so))
                self.attach_node(node, parent, content_tree);
                state["diem"] = node;
                return
            append_text_node(text)

        for block in iter_blocks_recursive(doc):
            if isinstance(block, docx.text.paragraph.Paragraph):
                text = collapse_ws(block.text)
                if not text: continue
                if self.config.stop_at_attachment_marker and re.search(r"FILE\s+ĐƯỢC\s+ĐÍNH\s+KÈM\s+THEO\s+VĂN\s+BẢN",
                                                                       text, re.IGNORECASE):
                    metadata["attachments_marker_found"] = True;
                    flush_quote();
                    break
                if self.config.stop_at_appendix and not ps["is_preamble"] and re.match(
                        r"^phụ\s*lục(\s+[IVXLCDM\d\w]+)?$", text, re.IGNORECASE):
                    metadata["appendix_marker_found"] = True;
                    flush_quote();
                    break
                if "Nơi nhận:" in text: ps["is_footer"] = True
                if ps["is_preamble"]:
                    if self.doc_type_pattern.match(text.upper()):
                        metadata["document_type"] = text;
                        metadata["document_title"] = text;
                        ps["collecting_title"] = True;
                        continue
                    if ps["collecting_title"]:
                        if text.lower().startswith("căn cứ"):
                            ps["collecting_title"] = False;
                            metadata["preamble"].append(text)
                        else:
                            metadata["document_title"] = collapse_ws(metadata.get("document_title", "") + " " + text)
                        continue
                is_bold_start = text.lower().startswith(("phần", "chương", "điều")) and any(
                    (run.text or "").strip() and run.bold for run in block.runs)
                if ps["is_preamble"] and (
                        self.pattern_phan.match(text) or self.pattern_chuong.match(text) or self.pattern_dieu.match(
                    text) or is_bold_start): ps["is_preamble"] = False
                if ps["is_preamble"]:
                    if "CỘNG HÒA" not in text.upper() and not re.search(r"\bSố\s*:", text, re.IGNORECASE): metadata[
                        "preamble"].append(text)
                    continue
                if ps["is_footer"]:
                    footer_data.append({"type": "text", "content": text});
                    continue
                process_text(text)
            elif isinstance(block, docx.table.Table):
                html = html_tables[table_idx] if table_idx < len(
                    html_tables) else "<table border='1'><tr><td>Lỗi Mammoth</td></tr></table>";
                table_idx += 1
                if block == doc.tables[0]: continue
                sig = False
                for r in block.rows:
                    if not r.cells: continue
                    left, right = collapse_ws(r.cells[0].text).upper(), collapse_ws(r.cells[-1].text).upper()
                    if "NƠI NHẬN" in left or (
                            re.match(r"^(CHỦ TỊCH|KT\.\s*BỘ TRƯỞNG|BỘ TRƯỞNG|TM\.\s*CHÍNH PHỦ|THỦ TƯỚNG|TM\.\s*ỦY BAN)",
                                     right) and not ps["is_preamble"]):
                        sig = True;
                        ps["is_footer"] = True;
                        break
                if sig:
                    footer_data.append({"type": "signature", "html": html});
                    break
                flush_quote();
                ps["pending_title_node"] = None
                if ps["is_footer"]:
                    footer_data.append({"type": "table", "html": html});
                    continue
                if ps["is_preamble"]: continue
                parent = self.current_parent(state);
                parent_id = parent["id"] if parent else document_id
                siblings = parent.get("children", []) if parent else content_tree
                idx = 1 + sum(1 for c in siblings if c.get("type") == "table")
                table_id = unique_node_id(f"{parent_id}.table_{idx}")
                rows = normalize_html_table(html) if self.config.normalize_tables else []
                node = self.make_node(document_id=document_id, node_type="table", no=str(idx), content=f"Bảng {idx}",
                                      parent=parent, path_stack=self.path_stack_from_state(metadata, state),
                                      node_id=table_id,
                                      extra={"html": html, "normalized_rows": rows, "content_hash": md5_text(html)})
                self.attach_node(node, parent, content_tree)
                tables_flat.append({"table_id": table_id, "document_id": document_id, "parent_id": parent_id,
                                    "path_text": node["path_text"], "html": html, "normalized_rows": rows,
                                    "row_count": len(rows), "col_count": max([len(r) for r in rows], default=0)})
        flush_quote()
        tree = {"metadata": metadata, "body": content_tree, "footer": footer_data}
        units = list(self.flatten_units(tree));
        refs = list(self.flatten_ref_mentions(units));
        amends = list(self.flatten_amendment_mentions(units))
        tree_path = doc_dir / f"{out_name}.tree.json";
        units_path = doc_dir / "units.jsonl";
        tables_path = doc_dir / "tables.jsonl";
        ref_path = doc_dir / "ref_mentions.jsonl";
        amend_path = doc_dir / "amendment_mentions.jsonl"
        if self.config.write_tree_json: write_json(tree_path, tree)
        if self.config.write_units_jsonl: write_jsonl(units_path, units)
        if self.config.write_tables_jsonl: write_jsonl(tables_path, tables_flat)
        if self.config.write_ref_mentions_jsonl: write_jsonl(ref_path, refs)
        if self.config.write_amendment_mentions_jsonl: write_jsonl(amend_path, amends)
        return ParseResult(doc_dir, tree_path, units_path, tables_path, ref_path, amend_path, document_id, len(units),
                           len(tables_flat), len(refs), len(amends))

    def flatten_units(self, tree: Dict[str, Any]) -> Iterator[Dict[str, Any]]:
        md = tree.get("metadata", {})
        doc_fields = {k: md.get(k) for k in
                      ["document_id", "document_number", "document_type", "document_title", "issue_date",
                       "issuing_agency", "source_file"]}

        def walk(node: Dict[str, Any]) -> Iterator[Dict[str, Any]]:
            unit = {k: v for k, v in node.items() if k != "children"};
            unit.update(doc_fields);
            unit["text_for_embedding"] = self.build_text_for_embedding(unit);
            yield unit
            for child in node.get("children", []) or []: yield from walk(child)

        for top in tree.get("body", []) or []: yield from walk(top)

    @staticmethod
    def build_text_for_embedding(unit: Dict[str, Any]) -> str:
        parts = [f"Văn bản: {unit.get('document_number') or unit.get('document_id')}",
                 f"Tên văn bản: {unit.get('document_title') or ''}", f"Đường dẫn: {unit.get('path_text') or ''}"]
        if unit.get("type") == "table":
            rows = unit.get("normalized_rows") or []
            parts.append("Nội dung bảng:\n" + "\n".join([" | ".join(r) for r in rows[:8]]))
        else:
            parts.append(f"Nội dung: {unit.get('content') or ''}")
        return "\n".join([p for p in parts if p.strip()])

    @staticmethod
    def flatten_ref_mentions(units: Iterable[Dict[str, Any]]) -> Iterator[Dict[str, Any]]:
        for unit in units:
            for ref in unit.get("ref_mentions", []) or []:
                yield {"source_unit_id": unit.get("id"), "document_id": unit.get("document_id"),
                       "document_number": unit.get("document_number"), "source_path_text": unit.get("path_text"),
                       "source_text": unit.get("content"), **ref}

    @staticmethod
    def flatten_amendment_mentions(units: Iterable[Dict[str, Any]]) -> Iterator[Dict[str, Any]]:
        for unit in units:
            for item in unit.get("amendment_mentions", []) or []:
                yield {"source_unit_id": unit.get("id"), "document_id": unit.get("document_id"),
                       "document_number": unit.get("document_number"), "source_path_text": unit.get("path_text"),
                       "source_text": unit.get("content"), **item}
