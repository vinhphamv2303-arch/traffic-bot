"""
Download missing Vietnamese legal documents from LuatVietnam as official Word/ZIP files.

V4 changes:
- Resolve the real LuatVietnam detail URL for each document_number.
- On the detail page, explicitly opens the "Tải về" tab/panel before looking for files.
- Downloads the Vietnamese DOC/DOCX file when there is one main Word file.
- If the page has multiple Vietnamese Word files/attachments and a "Tải tất cả" button, downloads the ZIP instead.
- Optional fallback can still reconstruct DOCX from public page text if official download fails.

Install:
  pip install pandas beautifulsoup4 lxml python-docx playwright tqdm
  python -m playwright install chromium

Login once:
  python download_luatvietnam_missing_v4.py --csv "...effectivity_index.csv" --login

Test one:
  python download_luatvietnam_missing_v4.py --csv "...effectivity_index.csv" --only "03/2018/TT-BGTVT" --out "...\\luatvietnam_official" --report "...\\download_report.csv" --headful --debug

Run all missing:
  python download_luatvietnam_missing_v4.py --csv "...effectivity_index.csv" --out "...\\luatvietnam_official" --report "...\\download_report.csv" --headful --debug --fallback-public-text
"""

from __future__ import annotations

import argparse
import csv
import re
import time
from dataclasses import dataclass
from pathlib import Path
from typing import Optional
from urllib.parse import quote_plus, urljoin, urlparse, parse_qs, unquote

import pandas as pd
from bs4 import BeautifulSoup
from docx import Document
from playwright.sync_api import sync_playwright, TimeoutError as PlaywrightTimeoutError
from tqdm import tqdm


BROWSER_PROFILE = Path("../.luatvietnam_browser_profile")
DEFAULT_START_URL = "https://luatvietnam.vn/xuat-nhap-khau/thong-tu-03-2018-tt-bgtvt-bo-giao-thong-van-tai-158515-d1.html"
SEARCH_URL = "https://luatvietnam.vn/van-ban/tim-van-ban.html"


@dataclass
class DownloadResult:
    document_id: str
    document_number: str
    url: Optional[str]
    status: str
    output_path: Optional[str]
    message: str


def safe_filename(s: str) -> str:
    s = str(s).lower().strip()
    s = s.replace("/", "_")
    s = re.sub(r"[^a-z0-9_\-.]+", "_", s)
    s = re.sub(r"_+", "_", s).strip("_")
    return s or "document"


def normalize_doc_number(s: str) -> str:
    s = str(s).upper().strip()
    s = s.replace(" ", "")
    s = s.replace("Đ", "D")
    return s


def doc_tokens(document_number: str) -> list[str]:
    return re.findall(r"[A-ZĐ0-9]+", normalize_doc_number(document_number))


def unwrap_search_url(href: str) -> str:
    if not href:
        return ""
    if href.startswith("/"):
        return urljoin("https://luatvietnam.vn", href)
    if "luatvietnam.vn" in href and href.startswith("http"):
        return href
    parsed = urlparse(href)
    qs = parse_qs(parsed.query)
    for key in ("u", "url", "q"):
        if key in qs:
            candidate = unquote(qs[key][0])
            if candidate.startswith("http") and "luatvietnam.vn" in candidate:
                return candidate
    return href


def is_detail_url(url: str) -> bool:
    return (
        url.startswith("http")
        and "luatvietnam.vn" in url
        and "english.luatvietnam.vn" not in url
        and "-d1.html" in url
    )


def score_candidate(url: str, text: str, document_number: str) -> int:
    target = normalize_doc_number(document_number)
    joined = normalize_doc_number((text or "") + " " + (url or ""))
    score = 0
    if target in joined:
        score += 100
    for t in doc_tokens(document_number):
        if t and t in joined:
            score += 10
    # Prefer detail pages and non-news pages.
    if "-d1.html" in url:
        score += 20
    if "-article.html" in url:
        score -= 20
    return score


def collect_luatvietnam_candidates(page, document_number: str) -> list[tuple[int, str, str]]:
    """Collect candidate detail URLs from current browser page."""
    anchors = page.evaluate(
        r"""() => Array.from(document.querySelectorAll('a[href]')).map(a => ({
            href: a.href || a.getAttribute('href') || '',
            text: (a.innerText || a.textContent || a.title || '').trim()
        }))"""
    )
    candidates: list[tuple[int, str, str]] = []
    seen = set()
    for a in anchors:
        href = unwrap_search_url(a.get("href", ""))
        if not is_detail_url(href):
            continue
        href = href.split("#")[0]
        if href in seen:
            continue
        seen.add(href)
        text = a.get("text", "")
        sc = score_candidate(href, text, document_number)
        candidates.append((sc, href, text))
    candidates.sort(key=lambda x: x[0], reverse=True)
    return candidates


def search_luatvietnam_in_site(page, document_number: str, debug: bool = False) -> Optional[str]:
    """Search LuatVietnam's own document search page.

    Important: the working document-search URL is /van-ban/tim-van-ban.html,
    not the generic /tim-van-ban.html page.  Query parameter observed on the
    site is `keywords`, so try direct URL search first, then the visible box.
    """
    query_variants = []
    dn = str(document_number).strip()
    query_variants.append(dn)
    query_variants.append(dn.replace("/", " "))
    query_variants.append(dn.replace("/", "-"))

    # 1) Direct URL search variants. This is more stable than filling the UI.
    for q in dict.fromkeys(query_variants):
        for param in ("keywords", "keyword", "q"):
            url = SEARCH_URL + f"?{param}=" + quote_plus(q)
            try:
                if debug:
                    print("  Try LuatVietnam search URL:", url)
                page.goto(url, wait_until="domcontentloaded", timeout=60000)
                page.wait_for_timeout(2000)
                candidates = collect_luatvietnam_candidates(page, document_number)
                if debug:
                    print("  LuatVietnam candidates:", candidates[:5])
                if candidates and candidates[0][0] >= 40:
                    return candidates[0][1]
            except Exception as e:
                if debug:
                    print("  LuatVietnam direct search failed:", e)

    # 2) UI fallback. Site HTML/React can change, so keep selectors broad.
    try:
        page.goto(SEARCH_URL, wait_until="domcontentloaded", timeout=60000)
        page.wait_for_timeout(1200)
    except Exception:
        return None

    selectors = [
        'input[name="keywords"]',
        'input[placeholder*="Tìm"]',
        'input[placeholder*="tìm"]',
        'input[type="search"]',
        'input[type="text"]',
        'textarea',
    ]
    for sel in selectors:
        try:
            count = page.locator(sel).count()
            for i in range(min(count, 10)):
                item = page.locator(sel).nth(i)
                if item.is_visible(timeout=1000):
                    item.click(timeout=3000)
                    item.fill(document_number, timeout=5000)
                    item.press("Enter", timeout=5000)
                    for _ in range(12):
                        page.wait_for_timeout(700)
                        candidates = collect_luatvietnam_candidates(page, document_number)
                        if debug:
                            print("  LuatVietnam UI candidates:", candidates[:5])
                        if candidates and candidates[0][0] >= 40:
                            return candidates[0][1]
        except Exception:
            continue

    return None


def search_bing_in_browser(page, document_number: str, debug: bool = False) -> Optional[str]:
    """Fallback search using Bing in the browser, not requests, to reduce bot-blocking problems."""
    queries = [
        f'site:luatvietnam.vn "{document_number}"',
        f'"{document_number}" "luatvietnam.vn"',
    ]
    for q in queries:
        page.goto("https://www.bing.com/search?q=" + quote_plus(q), wait_until="domcontentloaded", timeout=60000)
        page.wait_for_timeout(1800)
        candidates = collect_luatvietnam_candidates(page, document_number)
        if debug and candidates[:3]:
            print("  Bing candidates:", candidates[:3])
        if candidates and candidates[0][0] >= 50:
            return candidates[0][1]
    return None


def find_luatvietnam_url(page, document_number: str, debug: bool = False) -> Optional[str]:
    url = search_luatvietnam_in_site(page, document_number, debug=debug)
    if url:
        return url
    return search_bing_in_browser(page, document_number, debug=debug)


def ensure_login() -> None:
    with sync_playwright() as p:
        ctx = p.chromium.launch_persistent_context(
            str(BROWSER_PROFILE), headless=False, accept_downloads=True,
            viewport={"width": 1400, "height": 900},
        )
        page = ctx.new_page()
        page.goto(DEFAULT_START_URL, wait_until="domcontentloaded", timeout=60000)
        print("\nBrowser opened. Log in to LuatVietnam in that browser window.")
        input("After login succeeds, press Enter here to save the session and continue...")
        ctx.close()


def _click_text_like(page, wanted: str, exact: bool = True) -> bool:
    """Click the first visible clickable element whose text matches wanted."""
    js = r"""
    ({wanted, exact}) => {
      const norm = s => (s || '').replace(/\s+/g, ' ').trim().toLowerCase();
      const w = norm(wanted);
      const isVisible = el => {
        const r = el.getBoundingClientRect();
        const st = window.getComputedStyle(el);
        return r.width > 0 && r.height > 0 && st.visibility !== 'hidden' && st.display !== 'none';
      };
      const matches = el => {
        const t = norm(el.innerText || el.textContent || el.getAttribute('title') || el.getAttribute('aria-label') || '');
        return exact ? t === w : t.includes(w);
      };
      const clickable = el => {
        let cur = el;
        for (let i = 0; i < 8 && cur; i++, cur = cur.parentElement) {
          const tag = (cur.tagName || '').toLowerCase();
          const role = cur.getAttribute && cur.getAttribute('role');
          if (tag === 'a' || tag === 'button' || role === 'button' || cur.onclick) return cur;
        }
        return el;
      };
      const nodes = Array.from(document.querySelectorAll('a,button,span,div,li,label'));
      for (const n of nodes) {
        if (!isVisible(n)) continue;
        if (!matches(n)) continue;
        const c = clickable(n);
        c.scrollIntoView({block:'center'});
        c.click();
        return true;
      }
      return false;
    }
    """
    return bool(page.evaluate(js, {"wanted": wanted, "exact": exact}))


def open_download_panel(page, debug: bool = False) -> bool:
    """Open the LuatVietnam 'Tải về' tab/panel on a detail page."""
    # First try the visible tab text. On this site there can be several text nodes,
    # so JS walks up to the actual clickable tab/button.
    for text, exact in [("Tải về", True), ("Tải", False)]:
        try:
            if _click_text_like(page, text, exact=exact):
                page.wait_for_timeout(1200)
                if debug:
                    print("  Opened download panel by text:", text)
                return True
        except Exception:
            pass

    # URL hash fallback sometimes works after the page has rendered the tabs.
    try:
        page.evaluate("() => { location.hash = 'download'; window.dispatchEvent(new HashChangeEvent('hashchange')); }")
        page.wait_for_timeout(1000)
    except Exception:
        pass
    return False


def count_visible_vietnamese_word_items(page) -> int:
    """Count visible Vietnamese DOC/DOCX items in the download panel."""
    js = r"""
    () => {
      const norm = s => (s || '').replace(/\s+/g, ' ').trim().toLowerCase();
      const isVisible = el => {
        const r = el.getBoundingClientRect();
        const st = window.getComputedStyle(el);
        return r.width > 0 && r.height > 0 && st.visibility !== 'hidden' && st.display !== 'none';
      };
      const bad = ['văn bản tiếng anh', 'bản dịch tham khảo', 'circular ', 'english'];
      const nodes = Array.from(document.querySelectorAll('a,button,li,div,span,label'));
      const texts = [];
      for (const n of nodes) {
        if (!isVisible(n)) continue;
        const t = norm(n.innerText || n.textContent || n.getAttribute('title') || n.getAttribute('href') || '');
        if (!t) continue;
        if (!(t.includes('doc') || t.includes('docx') || t.includes('bản word') || t.includes('word'))) continue;
        if (bad.some(b => t.includes(b))) continue;
        // Avoid counting large containers that include both Vietnamese and English sections.
        if (t.length > 220) continue;
        texts.push(t);
      }
      return Array.from(new Set(texts)).length;
    }
    """
    try:
        return int(page.evaluate(js) or 0)
    except Exception:
        return 0


def click_download_all(page, debug: bool = False) -> bool:
    """Click 'Tải tất cả' if visible. This usually downloads a ZIP of attachments."""
    try:
        ok = _click_text_like(page, "Tải tất cả", exact=False)
        if ok and debug:
            print("  Clicked Tải tất cả")
        return ok
    except Exception:
        return False


def click_vietnamese_word_download(page, debug: bool = False) -> bool:
    """Click the first visible Vietnamese DOC/DOCX/Bản Word item, excluding English/translation items."""
    js = r"""
    () => {
      const norm = s => (s || '').replace(/\s+/g, ' ').trim().toLowerCase();
      const isVisible = el => {
        const r = el.getBoundingClientRect();
        const st = window.getComputedStyle(el);
        return r.width > 0 && r.height > 0 && st.visibility !== 'hidden' && st.display !== 'none';
      };
      const isWord = s => s.includes('doc') || s.includes('docx') || s.includes('bản word') || s.includes('word');
      const isBad = s => s.includes('văn bản tiếng anh') || s.includes('bản dịch tham khảo') || s.includes('circular ') || s.includes('english');
      const clickable = el => {
        let cur = el;
        for (let i = 0; i < 8 && cur; i++, cur = cur.parentElement) {
          const tag = (cur.tagName || '').toLowerCase();
          const role = cur.getAttribute && cur.getAttribute('role');
          if (tag === 'a' || tag === 'button' || role === 'button' || cur.onclick) return cur;
        }
        return el;
      };

      // Prefer real links/buttons over broad containers.
      const selectors = ['a[href]', 'button', 'label', 'span', 'div'];
      for (const sel of selectors) {
        const nodes = Array.from(document.querySelectorAll(sel));
        for (const n of nodes) {
          if (!isVisible(n)) continue;
          const t = norm(n.innerText || n.textContent || n.getAttribute('title') || n.getAttribute('aria-label') || n.getAttribute('href') || '');
          if (!t || !isWord(t) || isBad(t)) continue;
          if (t.length > 260) continue;
          const c = clickable(n);
          c.scrollIntoView({block:'center'});
          c.click();
          return true;
        }
      }
      return false;
    }
    """
    ok = bool(page.evaluate(js))
    if ok and debug:
        print("  Clicked Vietnamese Word item")
    return ok


def _target_path_from_download(output_base: Path, suggested_filename: str) -> Path:
    """Preserve the real downloaded extension (.doc, .docx, .zip, ...)."""
    name = suggested_filename or ""
    suffix = Path(name).suffix.lower()
    if suffix not in {".doc", ".docx", ".zip", ".pdf", ".rar"}:
        suffix = ".docx"
    return output_base.with_suffix(suffix)


def download_official_word(page, url: str, output_base: Path, debug: bool = False) -> tuple[bool, Optional[Path], str]:
    """Download official Vietnamese Word file or ZIP from LuatVietnam.

    The site does not show direct DOC links until the 'Tải về' tab is opened.
    If multiple Vietnamese DOC items exist and a 'Tải tất cả' button is present,
    we download the ZIP; otherwise we click the first Vietnamese Word file.
    """
    page.goto(url, wait_until="domcontentloaded", timeout=60000)
    page.wait_for_timeout(1500)
    open_download_panel(page, debug=debug)
    page.wait_for_timeout(1000)

    word_count = count_visible_vietnamese_word_items(page)
    if debug:
        print("  Visible Vietnamese Word items:", word_count)

    try:
        with page.expect_download(timeout=20000) as download_info:
            clicked = False
            if word_count >= 2:
                clicked = click_download_all(page, debug=debug)
            if not clicked:
                clicked = click_vietnamese_word_download(page, debug=debug)
            if not clicked:
                return False, None, "Đã mở tab Tải về nhưng không tìm thấy item DOC/Word tiếng Việt hoặc nút Tải tất cả"

        download = download_info.value
        final_path = _target_path_from_download(output_base, download.suggested_filename)
        tmp = final_path.with_suffix(final_path.suffix + ".part")
        final_path.parent.mkdir(parents=True, exist_ok=True)
        download.save_as(str(tmp))
        tmp.replace(final_path)
        return True, final_path, "OK: tải file gốc " + final_path.suffix

    except PlaywrightTimeoutError:
        try:
            text = page.locator("body").inner_text(timeout=5000)[:3000]
        except Exception:
            text = ""
        if "Đăng nhập" in text or "tài khoản thành viên" in text or "thành viên" in text:
            return False, None, "Click tải nhưng không phát sinh download: kiểm tra tài khoản đã đăng nhập và có quyền tải file thành viên"
        return False, None, "Click tải nhưng không phát sinh download; có thể còn popup/captcha hoặc endpoint đã đổi"
    except Exception as e:
        return False, None, f"Lỗi khi tải file Word/ZIP gốc: {e}"

def fallback_public_text_to_docx(page, url: str, output_path: Path) -> tuple[bool, str]:
    """Create a DOCX from public page text. This is not the official Word file."""
    page.goto(url, wait_until="domcontentloaded", timeout=60000)
    html = page.content()
    soup = BeautifulSoup(html, "lxml")

    for tag in soup(["script", "style", "noscript", "svg", "header", "footer", "nav", "aside"]):
        tag.decompose()

    title = soup.find("h1")
    doc = Document()
    if title:
        doc.add_heading(title.get_text(" ", strip=True), level=1)

    body_text = soup.get_text("\n", strip=True)

    # Start near the actual legal text, not page navigation.
    start_markers = [
        "BỘ ", "CHÍNH PHỦ", "QUỐC HỘI", "THỦ TƯỚNG CHÍNH PHỦ",
        "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM",
        "THÔNG TƯ", "NGHỊ ĐỊNH", "QUYẾT ĐỊNH", "LUẬT ",
    ]
    starts = [body_text.find(m) for m in start_markers if body_text.find(m) >= 0]
    if starts:
        body_text = body_text[min(starts):]

    stop_markers = ["Bạn chưa Đăng nhập", "Văn bản liên quan", "văn bản cùng lĩnh vực", "Tin liên quan", "Đánh giá bài viết"]
    stops = [body_text.find(m) for m in stop_markers if body_text.find(m) > 0]
    if stops:
        body_text = body_text[:min(stops)]

    lines = []
    for ln in body_text.splitlines():
        ln = re.sub(r"\s+", " ", ln).strip()
        if not ln:
            continue
        if ln in {"Đang theo dõi", "Đã biết", "Mục lục", "Tải về"}:
            continue
        if lines and lines[-1] == ln:
            continue
        lines.append(ln)

    for ln in lines:
        # Keep it simple; formatting is less important than text for RAG.
        doc.add_paragraph(ln)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return True, "Đã dựng DOCX từ nội dung public trên trang; không phải file Word gốc"


def read_existing_report(report_path: Path) -> set[str]:
    """Return document_ids already successfully processed, useful for --resume."""
    if not report_path.exists():
        return set()
    try:
        df = pd.read_csv(report_path)
        good = df[df["status"].isin(["downloaded_official", "fallback_public_text", "skipped"])]
        return set(good["document_id"].astype(str))
    except Exception:
        return set()


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--csv", required=True, help="Path to effectivity_index.csv")
    parser.add_argument("--out", default="data/raw_luatvietnam_docx", help="Output directory")
    parser.add_argument("--report", default="download_report.csv", help="Report CSV path")
    parser.add_argument("--only", default=None, help='Only download one document number, e.g. "03/2018/TT-BGTVT"')
    parser.add_argument("--login", action="store_true", help="Open browser to log in and save session first")
    parser.add_argument("--fallback-public-text", action="store_true", help="If official DOC fails, reconstruct DOCX from public HTML text")
    parser.add_argument("--headful", action="store_true", help="Run browser visibly during download")
    parser.add_argument("--sleep", type=float, default=0.8, help="Delay between documents")
    parser.add_argument("--debug", action="store_true", help="Print candidate URLs found during search")
    parser.add_argument("--resume", action="store_true", help="Skip document_ids already successful in existing report")
    args = parser.parse_args()

    if args.login:
        ensure_login()

    csv_path = Path(args.csv)
    if not csv_path.exists():
        raise FileNotFoundError(f"Không tìm thấy CSV: {csv_path}")

    df = pd.read_csv(csv_path)
    required_cols = {"document_id", "document_number", "in_corpus"}
    missing_cols = required_cols - set(df.columns)
    if missing_cols:
        raise ValueError(f"CSV thiếu cột: {sorted(missing_cols)}")

    missing = df[df["in_corpus"].astype(str).str.lower().eq("false")].copy()
    if args.only:
        missing = missing[missing["document_number"].astype(str).map(normalize_doc_number) == normalize_doc_number(args.only)]
    if args.resume:
        done_ids = read_existing_report(Path(args.report))
        missing = missing[~missing["document_id"].astype(str).isin(done_ids)]

    if missing.empty:
        print("Không có văn bản phù hợp để tải.")
        return

    out_dir = Path(args.out)
    out_dir.mkdir(parents=True, exist_ok=True)
    results: list[DownloadResult] = []

    with sync_playwright() as p:
        ctx = p.chromium.launch_persistent_context(
            str(BROWSER_PROFILE),
            headless=not args.headful,
            accept_downloads=True,
            viewport={"width": 1400, "height": 900},
        )
        page = ctx.new_page()

        for _, row in tqdm(missing.iterrows(), total=len(missing), desc="Downloading"):
            document_id = str(row["document_id"])
            document_number = str(row["document_number"])
            output_base = out_dir / safe_filename(document_id or document_number)

            existing = list(out_dir.glob(output_base.name + ".*"))
            existing = [p for p in existing if p.suffix.lower() in {".doc", ".docx", ".zip"} and p.stat().st_size > 0]
            if existing:
                results.append(DownloadResult(document_id, document_number, None, "skipped", str(existing[0]), "File đã tồn tại"))
                continue

            try:
                url = find_luatvietnam_url(page, document_number, debug=args.debug)
            except Exception as e:
                results.append(DownloadResult(document_id, document_number, None, "failed", None, f"Lỗi khi tìm URL: {e}"))
                continue

            if not url:
                results.append(DownloadResult(document_id, document_number, None, "failed", None, "Không tìm thấy URL LuatVietnam cho văn bản này"))
                continue

            ok, downloaded_path, msg = download_official_word(page, url, output_base, debug=args.debug)
            if ok and downloaded_path:
                status = "downloaded_zip" if downloaded_path.suffix.lower() == ".zip" else "downloaded_official"
                results.append(DownloadResult(document_id, document_number, url, status, str(downloaded_path), msg))
            elif args.fallback_public_text:
                fallback_path = output_base.with_suffix(".docx")
                ok2, msg2 = fallback_public_text_to_docx(page, url, fallback_path)
                results.append(
                    DownloadResult(
                        document_id,
                        document_number,
                        url,
                        "fallback_public_text" if ok2 else "failed",
                        str(fallback_path) if ok2 else None,
                        msg + " | " + msg2,
                    )
                )
            else:
                results.append(DownloadResult(document_id, document_number, url, "failed", None, msg))

            time.sleep(args.sleep)

        ctx.close()

    report_path = Path(args.report)
    report_path.parent.mkdir(parents=True, exist_ok=True)
    with report_path.open("w", newline="", encoding="utf-8-sig") as f:
        writer = csv.DictWriter(f, fieldnames=list(DownloadResult.__annotations__.keys()))
        writer.writeheader()
        for r in results:
            writer.writerow(r.__dict__)

    print(f"\nDone. Report: {report_path}")
    if results:
        print(pd.DataFrame([r.__dict__ for r in results])["status"].value_counts(dropna=False))


if __name__ == "__main__":
    main()
