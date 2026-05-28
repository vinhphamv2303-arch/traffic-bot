from __future__ import annotations

import argparse
import csv
import json
import logging
import re
import shutil
import sys
import tempfile
import unicodedata
import zipfile
from copy import deepcopy
from dataclasses import asdict, dataclass, field
from pathlib import Path, PurePosixPath
from typing import Dict, Iterable, List, Optional, Sequence, Tuple


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_RAW_ROOT = ROOT / "data" / "raw" / "luatvietnam_official"
DEFAULT_REPORT = ROOT / "data" / "raw" / "luatvietnam_download_report_v4.csv"
DEFAULT_OUTPUT_ROOT = ROOT / "data" / "dataset"
DEFAULT_LOG_FILE = ROOT / "data" / "raw" / "normalize_raw_dataset.log"
DEFAULT_SUMMARY = ROOT / "data" / "raw" / "normalize_raw_dataset_summary.json"

PARSER_MODULE_ROOT = ROOT / "data_preprocessing" / "legal_parser_modular"
if PARSER_MODULE_ROOT.exists():
    sys.path.insert(0, str(PARSER_MODULE_ROOT))

try:
    from legal_parser.common.doc_converter import convert_legacy_doc_file
except Exception:  # pragma: no cover - optional integration
    convert_legacy_doc_file = None


LEGAL_SOURCE_SUFFIXES = {".doc", ".docx", ".pdf", ".rtf", ".xls", ".xlsx"}
CONVERTIBLE_SUFFIXES = {".doc"}
WORD_LIKE_SUFFIXES = {".doc", ".docx", ".pdf", ".rtf"}
ZIP_SUFFIX = ".zip"

ATTACHMENT_KEYWORDS = (
    "phu luc",
    "phuluc",
    "mau so",
    "mauso",
    "bieu mau",
    "bieumau",
    "qcvn",
    "quy chuan",
    "tieu chuan",
    "danh muc",
    "bang",
    "appendix",
)

MAIN_KEYWORDS = (
    "van ban",
    "doc",
    "luat",
    "bo luat",
    "nghi dinh",
    "nghi quyet",
    "thong tu",
    "quyet dinh",
    "chi thi",
)


@dataclass
class SourceRecord:
    document_id: str
    document_number: str
    status: str
    source_path: Path
    from_report: bool = True


@dataclass
class OutputFile:
    role: str
    source_name: str
    output_name: str
    suffix: str
    size: int
    converted_from: Optional[str] = None
    conversion_status: Optional[str] = None
    conversion_message: Optional[str] = None
    split_from: Optional[str] = None


@dataclass
class EmbeddedHeading:
    start: int
    paragraph_index: int
    title: str
    key: str
    char_start: int = 0


@dataclass
class PackageResult:
    package_id: str
    document_id: str
    document_number: str
    source_path: str
    output_dir: str
    status: str
    message: str = ""
    files: List[OutputFile] = field(default_factory=list)
    warnings: List[str] = field(default_factory=list)


def strip_accents(text: str) -> str:
    text = text.replace("Đ", "D").replace("đ", "d")
    normalized = unicodedata.normalize("NFD", text)
    return "".join(ch for ch in normalized if unicodedata.category(ch) != "Mn")


def compact_text(text: str) -> str:
    return re.sub(r"[^a-z0-9]+", "", strip_accents(text).lower())


def normalized_text(text: str) -> str:
    text = strip_accents(text).lower()
    text = re.sub(r"[_\-]+", " ", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def clean_filename(name: str, *, max_len: int = 180) -> str:
    name = name.replace("\x00", " ")
    name = re.sub(r'[<>:"/\\|?*\x00-\x1f]+', " ", name)
    name = re.sub(r"\s+", " ", name).strip().rstrip(".")
    if not name:
        name = "file"
    if len(name) <= max_len:
        return name
    path = Path(name)
    suffix = path.suffix
    stem_max = max(12, max_len - len(suffix))
    return path.stem[:stem_max].rstrip(" .") + suffix


def safe_package_id(value: str) -> str:
    value = strip_accents(value)
    value = re.sub(r"[^A-Za-z0-9_]+", "_", value)
    value = re.sub(r"_+", "_", value).strip("_")
    return value.upper() or "UNKNOWN"


def package_id_from_document_id(document_id: str, document_number: str = "", source_name: str = "") -> str:
    candidates = [document_id, document_number.replace("/", "_").replace("-", "_"), source_name]
    for candidate in candidates:
        value = strip_accents(candidate or "").lower()
        parts = [p for p in re.split(r"[_\-\s/]+", value) if p]
        if len(parts) >= 3 and parts[0].isdigit() and re.fullmatch(r"\d{4}", parts[1]):
            suffix = "".join(p.upper() for p in parts[2:])
            return safe_package_id(f"{parts[0]}_{parts[1]}_{suffix}")

        match = re.search(r"(?<!\d)(\d{1,4})[^\d]{1,8}(\d{4})[^\w]+([a-z0-9_\- ]{2,40})", value)
        if match:
            number, year, tail = match.groups()
            tail_parts = [p for p in re.split(r"[_\-\s]+", tail) if p and not p.isdigit()]
            if tail_parts:
                suffix = "".join(part.upper() for part in tail_parts[:5])
                return safe_package_id(f"{number}_{year}_{suffix}")

    return safe_package_id(Path(source_name).stem or document_id or "unknown")


def is_attachment_name(name: str) -> bool:
    text = normalized_text(Path(name).stem)
    compact = compact_text(Path(name).stem)
    return any(keyword in text or keyword.replace(" ", "") in compact for keyword in ATTACHMENT_KEYWORDS)


def is_probable_main_name(name: str, package_id: str, document_number: str) -> bool:
    text = normalized_text(Path(name).stem)
    compact = compact_text(Path(name).stem)
    package_compact = compact_text(package_id)
    number_compact = compact_text(document_number)
    if package_compact and package_compact in compact:
        return True
    if number_compact and number_compact in compact:
        return True
    return any(keyword in text or keyword.replace(" ", "") in compact for keyword in MAIN_KEYWORDS)


def is_embedded_attachment_heading(text: str) -> bool:
    text = text.replace("\r", "").replace("\x07", "").strip()
    if not text or len(text) > 180:
        return False
    if re.match(r"^\s*mẫu\s+sổ\b", text.lower()):
        return False
    normalized = normalized_text(text)
    compact = compact_text(text)
    patterns = (
        r"^phu luc(\s+[ivxlcdm0-9a-z]+)?(\b|[\.:;-])",
        r"^mau so\s*[a-z0-9./-]+(\b|[\.:;-])",
        r"^qcvn\s*[0-9]",
        r"^quy chuan ky thuat quoc gia",
    )
    return any(re.search(pattern, normalized) for pattern in patterns) or compact.startswith("phuluc")


def clean_word_paragraph_text(text: str) -> str:
    return text.replace("\r", "").replace("\x07", "").strip()


def is_separator_text(text: str) -> bool:
    compact = re.sub(r"\s+", "", text)
    return bool(compact) and all(ch in "_-—=." for ch in compact)


def attachment_heading_key(text: str) -> str:
    normalized = normalized_text(text)
    match = re.match(r"^(phu luc)\s+([ivxlcdm0-9a-z]+)", normalized)
    if match:
        return f"{match.group(1)} {match.group(2)}"
    match = re.match(r"^(mau so)\s*([a-z0-9./-]+)", normalized)
    if match:
        return f"{match.group(1)} {match.group(2)}"
    return normalized


def unique_path(directory: Path, filename: str) -> Path:
    filename = clean_filename(filename)
    candidate = directory / filename
    if not candidate.exists():
        return candidate
    path = Path(filename)
    for index in range(2, 1000):
        candidate = directory / f"{path.stem} ({index}){path.suffix}"
        if not candidate.exists():
            return candidate
    raise RuntimeError(f"Cannot make a unique file name for {filename}")


def resolve_inside(path: Path, root: Path) -> bool:
    try:
        path.resolve().relative_to(root.resolve())
        return True
    except ValueError:
        return False


def remove_tree_safely(path: Path, root: Path) -> None:
    if not resolve_inside(path, root):
        raise RuntimeError(f"Refusing to remove outside output root: {path}")
    if path.exists():
        shutil.rmtree(path)


def load_report(report_path: Path, raw_root: Path) -> List[SourceRecord]:
    if not report_path.exists():
        return []

    records: List[SourceRecord] = []
    with report_path.open("r", encoding="utf-8-sig", newline="") as file:
        reader = csv.DictReader(file)
        for row in reader:
            output_path = (row.get("output_path") or "").strip()
            if not output_path:
                continue
            source_path = Path(output_path)
            if not source_path.exists():
                source_path = raw_root / Path(output_path).name
            if not source_path.exists():
                continue
            records.append(
                SourceRecord(
                    document_id=(row.get("document_id") or "").strip(),
                    document_number=(row.get("document_number") or "").strip(),
                    status=(row.get("status") or "").strip(),
                    source_path=source_path,
                    from_report=True,
                )
            )
    return records


def scan_raw_files(raw_root: Path, existing_sources: Iterable[Path]) -> List[SourceRecord]:
    existing = {str(path.resolve()).lower() for path in existing_sources}
    records: List[SourceRecord] = []
    if not raw_root.exists():
        return records
    for path in sorted(raw_root.iterdir()):
        if not path.is_file():
            continue
        if path.suffix.lower() not in LEGAL_SOURCE_SUFFIXES | {ZIP_SUFFIX}:
            continue
        if str(path.resolve()).lower() in existing:
            continue
        document_id = package_id_from_document_id("", "", path.name).lower()
        records.append(
            SourceRecord(
                document_id=document_id,
                document_number="",
                status="scanned_raw",
                source_path=path,
                from_report=False,
            )
        )
    return records


def safe_extract_zip(zip_path: Path, destination: Path, logger: logging.Logger) -> None:
    with zipfile.ZipFile(zip_path) as archive:
        for member in archive.infolist():
            if member.is_dir():
                continue
            raw_parts = PurePosixPath(member.filename.replace("\\", "/")).parts
            parts = [clean_filename(part) for part in raw_parts if part not in ("", ".", "..")]
            if not parts:
                continue
            target = destination.joinpath(*parts)
            if not resolve_inside(target, destination):
                logger.warning("Skip unsafe zip member %s in %s", member.filename, zip_path)
                continue
            target.parent.mkdir(parents=True, exist_ok=True)
            with archive.open(member) as source, target.open("wb") as output:
                shutil.copyfileobj(source, output)


def extract_zip_tree(zip_path: Path, destination: Path, logger: logging.Logger) -> None:
    safe_extract_zip(zip_path, destination, logger)
    processed = {zip_path.resolve()}
    while True:
        nested = [path for path in destination.rglob("*.zip") if path.resolve() not in processed]
        if not nested:
            break
        for nested_zip in nested:
            processed.add(nested_zip.resolve())
            nested_dir = nested_zip.with_suffix("")
            nested_dir.mkdir(parents=True, exist_ok=True)
            try:
                logger.info("Extract nested zip: %s", nested_zip)
                safe_extract_zip(nested_zip, nested_dir, logger)
            except zipfile.BadZipFile:
                logger.warning("Bad nested zip skipped: %s", nested_zip)


def collect_candidate_files(root: Path) -> List[Path]:
    files: List[Path] = []
    for path in sorted(root.rglob("*")):
        if not path.is_file():
            continue
        if path.name.startswith("~$"):
            continue
        if path.suffix.lower() in LEGAL_SOURCE_SUFFIXES:
            files.append(path)
    return files


def choose_main_file(files: Sequence[Path], package_id: str, document_number: str) -> Optional[Path]:
    if not files:
        return None
    if len(files) == 1:
        return files[0]

    def score(path: Path) -> Tuple[int, int]:
        suffix = path.suffix.lower()
        attachment = is_attachment_name(path.name)
        score_value = int(path.stat().st_size)
        if suffix == ".docx":
            score_value += 220_000_000
        elif suffix == ".doc":
            score_value += 200_000_000
        elif suffix in {".pdf", ".rtf"}:
            score_value += 60_000_000
        if suffix in {".xls", ".xlsx"}:
            score_value -= 100_000
        if attachment:
            score_value -= 500_000_000
        if is_probable_main_name(path.name, package_id, document_number):
            score_value += 1_000_000_000
        return score_value, -len(path.parts)

    return max(files, key=score)


def build_output_name(source: Path, role: str, package_id: str) -> str:
    suffix = source.suffix
    stem = source.stem.strip()
    if role == "main" and (stem.isdigit() or len(compact_text(stem)) <= 4):
        return f"{package_id} DOC{suffix}"
    return clean_filename(source.name)


def copy_package_files(
    candidate_files: Sequence[Path],
    main_file: Optional[Path],
    package_dir: Path,
    package_id: str,
) -> List[OutputFile]:
    outputs: List[OutputFile] = []
    for source in candidate_files:
        role = "main" if main_file is not None and source.resolve() == main_file.resolve() else "attachment"
        if role != "main" and not is_attachment_name(source.name):
            role = "unknown_attachment"
        output_name = build_output_name(source, role, package_id)
        target = unique_path(package_dir, output_name)
        target.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(source, target)
        outputs.append(
            OutputFile(
                role=role,
                source_name=str(source),
                output_name=target.name,
                suffix=target.suffix.lower(),
                size=target.stat().st_size,
            )
        )
    return outputs


def convert_docs(package_dir: Path, outputs: List[OutputFile], delete_doc: bool, logger: logging.Logger) -> None:
    if convert_legacy_doc_file is None:
        for output in outputs:
            if output.suffix == ".doc":
                output.conversion_status = "failed"
                output.conversion_message = "legal_parser doc converter is not importable"
        logger.warning("DOC converter is not importable; .doc files remain unchanged in %s", package_dir)
        return

    output_by_name: Dict[str, OutputFile] = {output.output_name: output for output in outputs}
    converter_logger = logging.getLogger("normalize_raw_dataset.doc_converter")
    converter_logger.handlers.clear()
    converter_logger.addHandler(logging.NullHandler())
    converter_logger.propagate = False
    converter_logger.setLevel(logging.CRITICAL + 1)

    for doc_path in sorted(package_dir.glob("*.doc")):
        if doc_path.name.startswith("~$"):
            continue
        logger.info("Convert DOC to DOCX: %s", doc_path)
        result = convert_legacy_doc_file(doc_path, delete_source=delete_doc, logger=converter_logger)
        output = output_by_name.get(doc_path.name)
        if output is None:
            continue
        output.conversion_status = result.status
        output.conversion_message = result.message
        if result.status in {"converted", "already_converted"} and result.target_path.exists():
            output.converted_from = output.output_name
            output.output_name = result.target_path.name
            output.suffix = result.target_path.suffix.lower()
            output.size = result.target_path.stat().st_size
            logger.info("Converted DOC to DOCX: %s -> %s", result.source_path.name, result.target_path.name)
        else:
            logger.warning("DOC conversion failed for %s: %s", doc_path.name, result.message)


def docx_content_children(document) -> List[object]:
    return [child for child in list(document.element.body) if not child.tag.endswith("}sectPr")]


def docx_block_text(block: object) -> str:
    texts: List[str] = []
    for node in block.iter():
        if node.tag.endswith("}t") and node.text:
            texts.append(node.text)
        elif node.tag.endswith("}tab"):
            texts.append("\t")
        elif node.tag.endswith("}br"):
            texts.append("\n")
    return clean_word_paragraph_text("".join(texts))


def build_embedded_attachment_title_from_blocks(blocks: Sequence[object], block_index: int) -> str:
    heading = docx_block_text(blocks[block_index])
    heading_key = attachment_heading_key(heading)
    parts = [heading]
    for next_index in range(block_index + 1, min(block_index + 8, len(blocks))):
        text = docx_block_text(blocks[next_index])
        if not text:
            continue
        if is_separator_text(text):
            continue
        normalized = normalized_text(text.strip("() /\\"))
        if normalized.startswith("ban hanh kem"):
            continue
        if attachment_heading_key(text) == heading_key:
            continue
        if is_embedded_attachment_heading(text):
            break
        if len(text) <= 220:
            parts.append(text)
        break
    return clean_filename(". ".join(parts))


def find_embedded_attachment_headings(document, min_main_chars: int) -> List[EmbeddedHeading]:
    headings: List[EmbeddedHeading] = []
    seen_starts: set[int] = set()
    blocks = docx_content_children(document)
    char_count = 0
    for index, block in enumerate(blocks):
        text = docx_block_text(block)
        if not is_embedded_attachment_heading(text):
            char_count += len(text) + 1
            continue
        if char_count < min_main_chars:
            char_count += len(text) + 1
            continue
        if index in seen_starts:
            char_count += len(text) + 1
            continue
        seen_starts.add(index)
        headings.append(
            EmbeddedHeading(
                start=index,
                paragraph_index=index,
                title=build_embedded_attachment_title_from_blocks(blocks, index),
                key=attachment_heading_key(text),
                char_start=char_count,
            )
        )
        char_count += len(text) + 1
    coalesced: List[EmbeddedHeading] = []
    for heading in headings:
        if coalesced and heading.key == coalesced[-1].key and heading.char_start - coalesced[-1].char_start < 2500:
            continue
        coalesced.append(heading)
    return coalesced


def save_docx_segment(
    blocks: Sequence[object],
    output_path: Path,
    start_index: int,
    end_index: int,
    *,
    source_path: Optional[Path] = None,
) -> None:
    from docx import Document

    output_path.parent.mkdir(parents=True, exist_ok=True)
    if output_path.exists():
        output_path.unlink()

    if source_path is not None:
        shutil.copy2(source_path, output_path)
        document = Document(output_path)
        body = document.element.body
        keep = set(range(start_index, end_index))
        for index, child in enumerate(docx_content_children(document)):
            if index not in keep:
                body.remove(child)
        document.save(output_path)
        return

    document = Document()
    body = document.element.body
    for child in docx_content_children(document):
        body.remove(child)
    for block in blocks[start_index:end_index]:
        body.append(deepcopy(block))
    document.save(output_path)


def split_embedded_attachments(
    package_dir: Path,
    outputs: List[OutputFile],
    *,
    min_main_chars: int,
    emit_attachments: bool,
    logger: logging.Logger,
) -> None:
    main_outputs = [output for output in outputs if output.role == "main" and output.suffix == ".docx"]
    if not main_outputs:
        return
    from docx import Document

    for output in main_outputs:
        main_path = package_dir / output.output_name
        if not main_path.exists():
            continue
        try:
            document = Document(main_path)
        except Exception as exc:
            logger.warning("Cannot inspect DOCX for embedded attachments %s: %s", main_path, exc)
            continue
        headings = find_embedded_attachment_headings(document, min_main_chars)
        if not headings:
            continue

        blocks = docx_content_children(document)
        if emit_attachments:
            logger.info("Split embedded attachments in %s: %s heading(s)", main_path.name, len(headings))
        else:
            logger.info("Trim embedded attachment tail from %s: %s heading(s)", main_path.name, len(headings))
        tmp_main = main_path.with_name(f"{main_path.stem}.__main_split_tmp{main_path.suffix}")
        save_docx_segment(blocks, tmp_main, 0, headings[0].start, source_path=main_path)

        new_outputs: List[OutputFile] = []
        if emit_attachments:
            for index, heading in enumerate(headings):
                next_start = headings[index + 1].start if index + 1 < len(headings) else len(blocks)
                if next_start <= heading.start:
                    continue
                attachment_name = clean_filename(f"{heading.title}.docx")
                attachment_path = unique_path(package_dir, attachment_name)
                save_docx_segment(blocks, attachment_path, heading.start, next_start, source_path=main_path)
                new_outputs.append(
                    OutputFile(
                        role="attachment",
                        source_name=f"embedded:{main_path.name}",
                        output_name=attachment_path.name,
                        suffix=".docx",
                        size=attachment_path.stat().st_size,
                        split_from=main_path.name,
                    )
                )

        main_path.unlink()
        tmp_main.rename(main_path)
        output.size = main_path.stat().st_size
        outputs.extend(new_outputs)


def write_manifest(package_dir: Path, result: PackageResult) -> None:
    manifest_path = package_dir / "dataset_manifest.json"
    with manifest_path.open("w", encoding="utf-8") as file:
        json.dump(asdict(result), file, ensure_ascii=False, indent=2)


def process_record(
    record: SourceRecord,
    output_root: Path,
    *,
    overwrite: bool,
    convert_doc: bool,
    delete_doc: bool,
    split_attachments: bool,
    min_main_chars_before_split: int,
    dry_run: bool,
    logger: logging.Logger,
) -> PackageResult:
    package_id = package_id_from_document_id(record.document_id, record.document_number, record.source_path.name)
    package_dir = output_root / package_id
    result = PackageResult(
        package_id=package_id,
        document_id=record.document_id,
        document_number=record.document_number,
        source_path=str(record.source_path),
        output_dir=str(package_dir),
        status="pending",
    )

    if package_dir.exists() and any(package_dir.iterdir()) and not overwrite:
        result.status = "skipped_existing"
        result.message = "Output package already exists. Use --overwrite to rebuild it."
        return result

    if dry_run:
        result.status = "dry_run"
        return result

    if overwrite:
        remove_tree_safely(package_dir, output_root)
    package_dir.mkdir(parents=True, exist_ok=True)

    with tempfile.TemporaryDirectory(prefix=f"{package_id}_", dir=str(output_root)) as temp_name:
        temp_dir = Path(temp_name)
        if record.source_path.suffix.lower() == ZIP_SUFFIX:
            logger.info("Extract zip source: %s", record.source_path)
            extract_zip_tree(record.source_path, temp_dir, logger)
            candidates = collect_candidate_files(temp_dir)
        else:
            candidates = [record.source_path]

        if not candidates:
            result.status = "failed"
            result.message = "No supported document files found."
            write_manifest(package_dir, result)
            return result

        main_file = choose_main_file(candidates, package_id, record.document_number)
        result.files = copy_package_files(candidates, main_file, package_dir, package_id)

    if convert_doc:
        convert_docs(package_dir, result.files, delete_doc=delete_doc, logger=logger)
    if split_attachments:
        has_external_attachments = any(
            item.role in {"attachment", "unknown_attachment"} and not item.split_from for item in result.files
        )
        split_embedded_attachments(
            package_dir,
            result.files,
            min_main_chars=min_main_chars_before_split,
            emit_attachments=not has_external_attachments,
            logger=logger,
        )

    main_count = sum(1 for item in result.files if item.role == "main")
    if main_count != 1:
        result.warnings.append(f"Expected exactly one main file, found {main_count}.")
    if any(item.suffix == ".doc" for item in result.files):
        result.warnings.append("Some .doc files remain because conversion failed or was disabled.")

    result.status = "created"
    result.message = f"Created {len(result.files)} file(s)."
    write_manifest(package_dir, result)
    return result


def setup_logging(log_file: Path, verbose: bool) -> logging.Logger:
    log_file.parent.mkdir(parents=True, exist_ok=True)
    logger = logging.getLogger("normalize_raw_dataset")
    logger.setLevel(logging.DEBUG if verbose else logging.INFO)
    logger.handlers.clear()

    formatter = logging.Formatter("%(asctime)s | %(levelname)s | %(message)s")
    console = logging.StreamHandler()
    console.setLevel(logging.DEBUG if verbose else logging.INFO)
    console.setFormatter(formatter)
    logger.addHandler(console)

    file_handler = logging.FileHandler(log_file, encoding="utf-8")
    file_handler.setLevel(logging.DEBUG)
    file_handler.setFormatter(formatter)
    logger.addHandler(file_handler)
    return logger


def build_records(raw_root: Path, report_path: Optional[Path]) -> List[SourceRecord]:
    report_records = load_report(report_path, raw_root) if report_path else []
    scanned = scan_raw_files(raw_root, [record.source_path for record in report_records])
    records = report_records + scanned
    seen: set[str] = set()
    unique: List[SourceRecord] = []
    for record in records:
        key = str(record.source_path.resolve()).lower()
        if key in seen:
            continue
        seen.add(key)
        unique.append(record)
    return unique


def write_summary(summary_path: Path, results: Sequence[PackageResult]) -> None:
    summary_path.parent.mkdir(parents=True, exist_ok=True)
    payload = {
        "total": len(results),
        "by_status": {},
        "results": [asdict(result) for result in results],
    }
    for result in results:
        payload["by_status"][result.status] = payload["by_status"].get(result.status, 0) + 1
    with summary_path.open("w", encoding="utf-8") as file:
        json.dump(payload, file, ensure_ascii=False, indent=2)


def parse_args(argv: Optional[Sequence[str]] = None) -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Normalize raw LuatVietnam files into the data/dataset package layout."
    )
    parser.add_argument("--raw-root", type=Path, default=DEFAULT_RAW_ROOT)
    parser.add_argument("--report", type=Path, default=DEFAULT_REPORT)
    parser.add_argument("--output-root", type=Path, default=DEFAULT_OUTPUT_ROOT)
    parser.add_argument("--summary", type=Path, default=DEFAULT_SUMMARY)
    parser.add_argument("--log-file", type=Path, default=DEFAULT_LOG_FILE)
    parser.add_argument("--limit", type=int, default=None)
    parser.add_argument(
        "--only",
        action="append",
        default=[],
        help="Process only records whose package id, document id, document number, or source name contains this value.",
    )
    parser.add_argument("--overwrite", action="store_true")
    parser.add_argument("--dry-run", action="store_true")
    parser.add_argument("--no-convert-doc", action="store_true")
    parser.add_argument(
        "--no-split-embedded-attachments",
        action="store_true",
        help="Disable Word-based splitting of appendix/form sections embedded in the main .docx.",
    )
    parser.add_argument(
        "--min-main-chars-before-split",
        type=int,
        default=2000,
        help="Ignore appendix/form headings too close to the start of the main document.",
    )
    parser.add_argument("--keep-doc", action="store_true", help="Keep copied .doc files after successful .docx conversion.")
    parser.add_argument("--verbose", action="store_true")
    return parser.parse_args(argv)


def main(argv: Optional[Sequence[str]] = None) -> int:
    args = parse_args(argv)
    logger = setup_logging(args.log_file, args.verbose)

    raw_root = args.raw_root.resolve()
    output_root = args.output_root.resolve()
    report = args.report.resolve() if args.report else None
    output_root.mkdir(parents=True, exist_ok=True)

    records = build_records(raw_root, report)
    if args.only:
        needles = [compact_text(value) for value in args.only if compact_text(value)]
        filtered: List[SourceRecord] = []
        for record in records:
            package_id = package_id_from_document_id(record.document_id, record.document_number, record.source_path.name)
            haystack = compact_text(
                " ".join([package_id, record.document_id, record.document_number, record.source_path.name])
            )
            if any(needle in haystack for needle in needles):
                filtered.append(record)
        records = filtered
    if args.limit is not None:
        records = records[: args.limit]

    logger.info("Raw root: %s", raw_root)
    logger.info("Output root: %s", output_root)
    logger.info("Records to process: %s", len(records))

    results: List[PackageResult] = []
    for index, record in enumerate(records, start=1):
        logger.info("[%s/%s] %s", index, len(records), record.source_path.name)
        try:
            result = process_record(
                record,
                output_root,
                overwrite=args.overwrite,
                convert_doc=not args.no_convert_doc,
                delete_doc=not args.keep_doc,
                split_attachments=not args.no_split_embedded_attachments,
                min_main_chars_before_split=args.min_main_chars_before_split,
                dry_run=args.dry_run,
                logger=logger,
            )
        except Exception as exc:
            package_id = package_id_from_document_id(record.document_id, record.document_number, record.source_path.name)
            logger.exception("Failed to process %s", record.source_path)
            result = PackageResult(
                package_id=package_id,
                document_id=record.document_id,
                document_number=record.document_number,
                source_path=str(record.source_path),
                output_dir=str(output_root / package_id),
                status="failed",
                message=str(exc),
            )
        logger.info("Result %s: %s", result.package_id, result.status)
        results.append(result)

    write_summary(args.summary, results)
    logger.info("Summary written to: %s", args.summary)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
